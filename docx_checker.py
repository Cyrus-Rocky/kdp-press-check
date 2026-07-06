"""KDP paperback-interior checks for Word (.docx) manuscripts.

Word doesn't store fixed page geometry the way a print-ready PDF does — there's
no rendered page count without actually paginating the text, and "bleed" only
exists once a file is exported to PDF. Where that matters, the checks here say
so explicitly rather than guessing a number and presenting it as exact. KDP can
auto-convert a Word file at upload time, but you don't see the converted result
until after you submit it, so checking before export still catches the things
that would otherwise surprise you.
"""
import io
from collections import Counter

from docx import Document
from PIL import Image

import classify
import content_quality
import frontmatter
import kdp_rules as rules

EMU_PER_INCH = 914400


def _emu_to_in(emu) -> float:
    return (emu or 0) / EMU_PER_INCH


def _first_page_text_docx(doc, max_fallback_paragraphs: int = 5) -> str:
    """Text up to the first explicit (manual) page break, or the first few
    non-empty paragraphs if no manual break is found — Word doesn't expose
    rendered page boundaries without actually paginating the document."""
    ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    collected = []
    non_empty_count = 0
    for p in doc.paragraphs:
        has_page_break = any(
            br.get(f"{ns}type") == "page"
            for run in p.runs
            for br in run._element.findall(f"{ns}br")
        )
        if p.text.strip():
            collected.append(p.text)
            non_empty_count += 1
        if has_page_break:
            break
        if non_empty_count >= max_fallback_paragraphs:
            break
    return "\n".join(collected)


def check_content_type_docx(doc) -> dict:
    full_text = "\n".join(p.text for p in doc.paragraphs)
    word_count = len(full_text.split())
    # Word has no fixed page count without rendering; estimate from character
    # volume against a typical paperback's ~280 words/page so the classifier
    # has a denominator. This estimate is for classification only — it is
    # never used to size margins.
    estimated_pages = max(1, round(word_count / 280)) if word_count else 1

    fonts_used = set()
    for p in doc.paragraphs:
        for run in p.runs:
            if run.font.name:
                fonts_used.add(run.font.name)
    mono_ratio = 1.0 if any("courier" in f.lower() or "mono" in f.lower()
                            for f in fonts_used) else 0.0

    result = classify.classify(estimated_pages, full_text, mono_ratio, 0.0)
    ok = result["kind"] == "book"
    out = {"title": "Document Type", "ok": ok, "summary": result["summary"]}
    if result.get("fix"):
        out["fix"] = result["fix"]
    out["detail"] = f"~{estimated_pages} page(s) estimated from word count. Classified as: {result['kind']}."
    return out


def check_trim_size_docx(doc) -> dict:
    section = doc.sections[0]
    w_in, h_in = _emu_to_in(section.page_width), _emu_to_in(section.page_height)
    match, dist = rules.closest_trim_size(w_in, h_in)
    ok = dist is not None and dist <= 0.1

    if ok:
        return {
            "title": "Trim Size", "ok": True,
            "summary": f"Your page size is set to {w_in:.2f}\" x {h_in:.2f}\" — a standard KDP trim size.",
            "detail": f"Matches KDP's {match[0]}\" x {match[1]}\" trim size exactly.",
        }
    return {
        "title": "Trim Size", "ok": False,
        "summary": f"Your page size is {w_in:.2f}\" x {h_in:.2f}\", which isn't a standard KDP trim size.",
        "fix": f"In Word: Layout > Size > More Paper Sizes, and set the page size to "
               f"{match[0]}\" x {match[1]}\" (the closest standard trim) or another size "
               f"from KDP's trim size list.",
        "detail": f"Measured {w_in:.2f}\" x {h_in:.2f}\". Closest standard size: {match[0]}\" x {match[1]}\".",
    }


def check_page_size_consistency_docx(doc) -> dict:
    sizes = {(round(s.page_width, -3), round(s.page_height, -3)) for s in doc.sections}
    ok = len(sizes) == 1
    if ok:
        return {"title": "Page Size Consistency", "ok": True,
                "summary": "All sections use the same page size.",
                "detail": "Single consistent page size found."}
    return {
        "title": "Page Size Consistency", "ok": False,
        "summary": f"This document has {len(sizes)} sections with different page sizes.",
        "fix": "Check each section break in your document (Layout > Breaks) and make sure "
               "every section uses the same page size before exporting.",
        "detail": f"Found {len(sizes)} distinct page sizes across {len(doc.sections)} section(s).",
    }


def check_margins_docx(doc) -> dict:
    section = doc.sections[0]
    left_in = _emu_to_in(section.left_margin) + _emu_to_in(section.gutter)
    right_in = _emu_to_in(section.right_margin)
    top_in = _emu_to_in(section.top_margin)
    bottom_in = _emu_to_in(section.bottom_margin)
    outer_min = rules.MIN_OUTSIDE_TOP_BOTTOM_MARGIN_IN
    inside_floor = rules.inside_margin_in(1)  # loosest tier, 0.375" — the absolute floor

    problems = []
    if left_in < inside_floor - rules.TOLERANCE_IN:
        problems.append(f"inside margin (left + gutter) is {left_in:.2f}\", below the "
                         f"absolute minimum of {inside_floor:.2f}\"")
    if right_in < outer_min - rules.TOLERANCE_IN:
        problems.append(f"right margin is {right_in:.2f}\", below the {outer_min:.2f}\" minimum")
    if top_in < outer_min - rules.TOLERANCE_IN:
        problems.append(f"top margin is {top_in:.2f}\", below the {outer_min:.2f}\" minimum")
    if bottom_in < outer_min - rules.TOLERANCE_IN:
        problems.append(f"bottom margin is {bottom_in:.2f}\", below the {outer_min:.2f}\" minimum")

    if problems:
        return {
            "title": "Margins", "ok": False,
            "summary": "Margin too tight: " + "; ".join(problems) + ".",
            "fix": "In Word: Layout > Margins > Custom Margins, and increase the margins "
                   f"listed above. Note: KDP's exact inside-margin requirement depends on "
                   f"your final printed page count ({inside_floor:.2f}\"-0.875\", longer "
                   f"books need more) — Word can't tell us that page count without "
                   f"rendering. Export to PDF and run it through Press Check for the precise "
                   f"number once your manuscript is final.",
            "detail": f"Left+gutter {left_in:.2f}\", right {right_in:.2f}\", top {top_in:.2f}\", "
                      f"bottom {bottom_in:.2f}\".",
        }

    return {
        "title": "Margins", "ok": True, "warning_only": True,
        "summary": f"Margins clear the absolute minimums "
                   f"(inside {inside_floor:.2f}\", outside/top/bottom {outer_min:.2f}\").",
        "fix": "Word can't tell us your final printed page count, and KDP's exact "
               "inside-margin requirement scales with it (0.375\"-0.875\"). Export to PDF "
               "once your manuscript is final and run it through Press Check again for the "
               "precise number.",
        "detail": f"Left+gutter {left_in:.2f}\", right {right_in:.2f}\", top {top_in:.2f}\", "
                  f"bottom {bottom_in:.2f}\".",
    }


def check_fonts_embedded_docx(doc) -> dict:
    settings_xml = doc.settings.element
    ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    embedded = settings_xml.find(f"{ns}embedTrueTypeFonts") is not None

    if embedded:
        return {"title": "Font Embedding", "ok": True,
                "summary": "This file is set to embed fonts.",
                "detail": "embedTrueTypeFonts is enabled in document settings."}
    return {
        "title": "Font Embedding", "ok": False, "warning_only": True,
        "summary": "Fonts aren't set to embed in this Word file (Word's default).",
        "fix": "In Word: File > Options > Save > check \"Embed fonts in the file\" before "
               "your final export. This mostly matters if you used a non-standard font — "
               "if you only used common fonts like Times New Roman or Calibri, KDP's "
               "converter will likely substitute correctly anyway.",
        "detail": "embedTrueTypeFonts not found in document settings.",
    }


def check_image_resolution_docx(doc) -> dict:
    low_res = []
    checked = 0
    for shape in doc.inline_shapes:
        try:
            blip = shape._inline.graphic.graphicData.pic.blipFill.blip
            rId = blip.embed
            image_part = doc.part.related_parts[rId]
            with Image.open(io.BytesIO(image_part.blob)) as img:
                pix_w, pix_h = img.size
        except Exception:
            continue

        disp_w_in = _emu_to_in(shape.width)
        disp_h_in = _emu_to_in(shape.height)
        if disp_w_in <= 0 or disp_h_in <= 0:
            continue
        checked += 1
        effective_dpi = min(pix_w / disp_w_in, pix_h / disp_h_in)
        if effective_dpi < rules.MIN_IMAGE_DPI:
            low_res.append(f"{effective_dpi:.0f} DPI (needs {rules.MIN_IMAGE_DPI}+)")

    if checked == 0:
        return {"title": "Image Resolution", "ok": True,
                "summary": "No inline images found to check.",
                "fix": None,
                "detail": "No inline images found. Note: images anchored as floating/"
                          "wrapped objects aren't checked in Word files — convert to PDF "
                          "for a complete image check."}
    if not low_res:
        return {
            "title": "Image Resolution", "ok": True,
            "summary": f"All {checked} inline image(s) meet the {rules.MIN_IMAGE_DPI} DPI minimum.",
            "detail": f"Checked {checked} inline image(s); all sharp enough to print. "
                      f"Floating/wrapped images aren't checked here — convert to PDF for "
                      f"a complete check.",
        }
    return {
        "title": "Image Resolution", "ok": False,
        "summary": f"{len(low_res)} of {checked} inline image(s) will print blurry.",
        "fix": f"Replace each one with its original high-resolution file, or re-export/"
               f"re-scan it at {rules.MIN_IMAGE_DPI} DPI or higher at the size you're "
               f"displaying it.",
        "detail": "\n".join(low_res[:10])
                  + ("\nNote: floating/wrapped images aren't checked here — convert to "
                     "PDF for a complete image check."),
    }


def check_font_consistency_docx(doc) -> dict:
    """Flag mixed font families/sizes in body text — a common artifact of
    pasting in text from another document or a web page, which brings its
    own formatting along and looks unprofessional once printed."""
    font_counts = Counter()
    size_counts = Counter()
    for p in doc.paragraphs:
        style_name = p.style.name if p.style else ""
        if style_name.startswith("Heading") or style_name.startswith("Title"):
            continue
        for run in p.runs:
            if not run.text.strip():
                continue
            if run.font.name:
                font_counts[run.font.name] += 1
            if run.font.size:
                size_counts[run.font.size.pt] += 1

    # A handful of stray runs (e.g. one italicized word) isn't worth flagging —
    # only surface it once a real second font/size shows up repeatedly.
    minority_threshold = 3
    problems = []
    if len(font_counts) > 1:
        minority = sum(font_counts.values()) - font_counts.most_common(1)[0][1]
        if minority >= minority_threshold:
            fonts_list = ", ".join(f"{f} ({n})" for f, n in font_counts.most_common())
            problems.append(f"body text mixes {len(font_counts)} font families: {fonts_list}")
    if len(size_counts) > 1:
        minority = sum(size_counts.values()) - size_counts.most_common(1)[0][1]
        if minority >= minority_threshold:
            sizes_list = ", ".join(f"{s:g}pt ({n})" for s, n in size_counts.most_common())
            problems.append(f"body text mixes {len(size_counts)} font sizes: {sizes_list}")

    if not problems:
        return {
            "title": "Font Consistency", "ok": True, "warning_only": True,
            "summary": "Body text uses a consistent font throughout.",
            "detail": "No mixed font-family or font-size overrides found in body paragraphs.",
        }
    return {
        "title": "Font Consistency", "ok": False, "warning_only": True,
        "summary": "Inconsistent formatting found: " + "; ".join(problems) + ".",
        "fix": "This usually happens when text is pasted in from another document or a web "
               "page and brings its own formatting along. Select the whole manuscript body, "
               "clear direct formatting (Ctrl+Spacebar in Word), and re-apply a single font "
               "and size throughout.",
        "detail": "; ".join(problems),
    }


def check_metadata_docx(doc) -> dict:
    title = (doc.core_properties.title or "").strip()
    if title:
        return {"title": "Metadata", "ok": True, "warning_only": True,
                "summary": f"Document title is set to \"{title}\".",
                "detail": f"Title property: \"{title}\"."}
    return {
        "title": "Metadata", "ok": False, "warning_only": True,
        "summary": "The document's title property is empty.",
        "fix": "Not required by KDP, but worth doing: File > Info > Properties > Title, "
               "and set it to match your book's title.",
        "detail": "Title property is empty.",
    }


def run_all_checks_docx(docx_path: str) -> dict:
    doc = Document(docx_path)
    results = [
        check_content_type_docx(doc),
        check_trim_size_docx(doc),
        check_page_size_consistency_docx(doc),
        check_margins_docx(doc),
        check_fonts_embedded_docx(doc),
        check_font_consistency_docx(doc),
        check_image_resolution_docx(doc),
        check_metadata_docx(doc),
    ]
    full_text = "\n".join(p.text for p in doc.paragraphs)
    headings = [p.text for p in doc.paragraphs
                if p.style and p.style.name and p.style.name.startswith("Heading") and p.text.strip()]
    results += content_quality.run(full_text, headings)

    doc_title = (doc.core_properties.title or "").strip() or None
    first_page_text = _first_page_text_docx(doc)
    results += frontmatter.run(full_text, first_page_text, doc_title)

    blocking_results = [r for r in results if not r.get("warning_only")]
    summary_ok = all(r["ok"] for r in blocking_results)
    issue_count = sum(1 for r in blocking_results if not r["ok"])
    advisory_issue_count = sum(1 for r in results if r.get("warning_only") and not r["ok"])
    word_count = len(full_text.split())
    estimated_pages = max(1, round(word_count / 280)) if word_count else 1
    return {
        "page_count": estimated_pages,
        "page_count_is_estimate": True,
        "results": results,
        "overall_ok": summary_ok,
        "issue_count": issue_count,
        "advisory_issue_count": advisory_issue_count,
    }
