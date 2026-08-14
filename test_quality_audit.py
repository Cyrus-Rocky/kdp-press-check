#!/usr/bin/env python3
"""Quality audit of KDP Press Check content quality checks"""

import sys
import io
from pathlib import Path

if sys.stdout.encoding != 'utf-8':
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

sys.path.insert(0, str(Path(__file__).parent))

try:
    from docx import Document
    import content_quality
    import page_breaks
    import frontmatter

    book_path = Path(r"C:\Users\HP\Downloads\Business Banking Basics EDITED.docx")

    if not book_path.exists():
        print("[ERROR] File not found: {}".format(book_path))
        sys.exit(1)

    print("[AUDIT] Quality Check: {}".format(book_path.name))
    print("=" * 70)

    doc = Document(str(book_path))
    full_text = "\n".join([p.text for p in doc.paragraphs])
    headings = [p.text for p in doc.paragraphs if p.style.name.startswith('Heading')]
    first_para = doc.paragraphs[0].text if doc.paragraphs else ""

    print("[OK] Loaded {} paragraphs".format(len(doc.paragraphs)))
    print("[OK] Found {} headings".format(len(headings)))
    print("[OK] Text length: {} characters".format(len(full_text)))
    print()

    print("[RUN] Quality Checks (content, structure, front/back matter)...")
    all_results = []
    all_results += content_quality.run(full_text, headings)
    all_results += page_breaks.run(full_text)
    all_results += frontmatter.run(full_text, first_para, None)

    blocking = [r for r in all_results if not r.get('ok') and not r.get('warning_only')]
    advisory = [r for r in all_results if r.get('warning_only') and not r.get('ok')]
    passed = [r for r in all_results if r.get('ok')]

    print()
    print("RESULTS:")
    print("=" * 70)

    print("\nPASSED ({0}):".format(len(passed)))
    for r in passed:
        print("  [PASS] {}".format(r['title']))
        if 'summary' in r and r['summary']:
            print("         {}".format(r['summary'][:75]))

    if advisory:
        print("\nADVISORY/QUALITY TIPS ({0}):".format(len(advisory)))
        for r in advisory:
            print("  [WARN] {}".format(r['title']))
            if 'summary' in r and r['summary']:
                print("         {}".format(r['summary'][:75]))
            if 'detail' in r and r['detail']:
                print("         Detail: {}".format(r['detail'][:90]))

    if blocking:
        print("\nBLOCKING ISSUES ({0}):".format(len(blocking)))
        for r in blocking:
            print("  [FAIL] {}".format(r['title']))
            if 'summary' in r and r['summary']:
                print("         {}".format(r['summary'][:75]))
            if 'detail' in r and r['detail']:
                print("         Detail: {}".format(r['detail'][:90]))

    print()
    print("=" * 70)
    print("ANALYSIS:")
    print("  Total Checks: {}".format(len(all_results)))
    print("  Passed:       {}".format(len(passed)))
    print("  Advisory:     {}".format(len(advisory)))
    print("  Blocking:     {}".format(len(blocking)))
    print()
    print("Quality Audit Assessment:")
    print("  - Repeated Words: {} issues".format(len([r for r in all_results if 'Repeated' in r.get('title', '')])))
    print("  - Spacing:  {} issues".format(len([r for r in all_results if 'Spacing' in r.get('title', '')])))
    print("  - Spelling: {} issues".format(len([r for r in all_results if 'Spelling' in r.get('title', '')])))
    print("  - Quotes:   {} issues".format(len([r for r in all_results if 'Quote' in r.get('title', '')])))
    print()

    if blocking:
        print("VERDICT: Content has BLOCKING ISSUES")
    elif advisory:
        print("VERDICT: Content is ACCEPTABLE but has quality suggestions")
    else:
        print("VERDICT: Content PASSES all quality checks")

except Exception as e:
    print("[ERROR] {}".format(e))
    import traceback
    traceback.print_exc()
    sys.exit(1)
